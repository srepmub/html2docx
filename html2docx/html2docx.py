import re
from html.parser import HTMLParser
from typing import Any, Dict, Iterator, List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from tinycss2 import parse_declaration_list
from tinycss2.ast import DimensionToken, IdentToken

from .image import image_size, load_image

WHITESPACE_RE = re.compile(r"\s+")

ALIGNMENTS = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def get_attr(attrs: List[Tuple[str, Optional[str]]], attr_name: str) -> str:
    value = next((val for name, val in attrs if name == attr_name), "")
    if value is None:
        raise AttributeError(attr_name)
    return value


def style_to_css(style: str) -> Iterator[Dict[str, Any]]:
    for declaration in parse_declaration_list(style):
        for value in declaration.value:
            if isinstance(value, DimensionToken):
                yield {
                    "name": declaration.lower_name,
                    "value": value.value,
                    "unit": value.lower_unit,
                }
            elif isinstance(value, IdentToken):
                yield {"name": declaration.lower_name, "value": value.lower_value}


def html_attrs_to_font_style(
    attrs: List[Tuple[str, Optional[str]]]
) -> List[Tuple[str, Any]]:
    """Return Font style names based on tag style attributes

    :param attrs:
        a list of attribute name plus attribute value tuples.
    :returns:
        a list of style names to be applied to Run Font property
    """
    styles = []
    style = get_attr(attrs, "style")
    for style_decl in style_to_css(style):
        if style_decl["name"] == "text-decoration":
            if style_decl["value"] == "underline":
                styles.append(("underline", True))
            elif style_decl["value"] == "line-through":
                styles.append(("strike", True))
    return styles


class HTML2Docx(HTMLParser):
    def __init__(self, title: str):
        super().__init__()
        self.doc = Document()
        self.doc.core_properties.title = title
        self.list_style: List[str] = []
        self.href = ""
        self.anchor = ""
        self._reset()

    def _reset(self) -> None:
        self.p: Optional[Paragraph] = None
        self.r: Optional[Run] = None

        # Formatting options
        self.pre = False
        self.style = False
        self.code = False
        self.table_cell: Optional[Any] = None
        self.tables: List[Tuple[Any, int, int]] = []
        self.alignment: Optional[int] = None
        self.padding_left: Optional[Pt] = None
        self.attrs: List[List[Tuple[str, Any]]] = []
        self.collapse_space = True

    def init_p(self, attrs: List[Tuple[str, Optional[str]]]) -> None:
        align = get_attr(attrs, "align")
        if align:
            self.alignment = ALIGNMENTS.get(align, WD_ALIGN_PARAGRAPH.LEFT)
        style = get_attr(attrs, "style")
        for style_decl in style_to_css(style):
            if style_decl["name"] == "text-align":
                self.alignment = ALIGNMENTS.get(
                    style_decl["value"], WD_ALIGN_PARAGRAPH.LEFT
                )
            elif style_decl["name"] == "padding-left" and style_decl["unit"] == "px":
                self.padding_left = Pt(style_decl["value"])

    def finish_p(self) -> None:
        if self.r is not None:
            self.r.text = self.r.text.rstrip()
        self._reset()

    def init_table(self, attrs: List[Tuple[str, Optional[str]]]) -> None:
        if self.table_cell is not None:
            table = self.table_cell.add_table(rows=0, cols=0)
        else:
            table = self.doc.add_table(rows=0, cols=0)
        self.tables.append((table, -1, -1))

    def finish_table(self) -> None:
        table = self.tables.pop()[0]
        section = self.doc.sections[0]
        page_width = section.page_width - section.left_margin - section.right_margin
        page_width = int(page_width * (0.5 ** len(self.tables)))
        for col in table.columns:
            col.width = page_width // len(table.columns)

    def init_tr(self) -> None:
        table, row, col = self.tables[-1]
        row += 1
        col = -1
        self.tables[-1] = (table, row, col)
        table.add_row()

    def init_tdth(self) -> None:
        table, row, col = self.tables[-1]
        col += 1
        self.tables[-1] = (table, row, col)
        if col >= len(table.columns):
            table.add_column(0)
        self.table_cell = table.cell(row, col)
        self.p = self.table_cell.paragraphs[0]
        self.r = None

    def init_run(self, attrs: List[Tuple[str, Any]]) -> None:
        self.attrs.append(attrs)
        if attrs:
            self.r = None

    def finish_run(self) -> None:
        attrs = self.attrs.pop()
        if attrs:
            self.r = None

    def add_text(self, data: str) -> None:
        if self.p is None:
            style = self.list_style[-1] if self.list_style else None
            self.p = self.doc.add_paragraph(style=style)
            if self.alignment is not None:
                self.p.alignment = self.alignment
        if self.padding_left:
            self.p.paragraph_format.left_indent = self.padding_left
        if self.r is None:
            self.r = self.p.add_run()
            for attrs in self.attrs:
                for font_attr, value in attrs:
                    setattr(self.r.font, font_attr, value)
        if self.href:
            self.add_hyperlink(self.href, data)
        elif self.anchor:
            self.add_bookmark(self.anchor, data)
        else:
            self.r.add_text(data)

    def add_hyperlink(self, href: str, text: str) -> None:
        if not href.startswith("#"):  # TODO external links
            if text.endswith(" "):
                text += href + " "
            else:
                text += " " + href
            if self.r:
                self.r.add_text(text)
            return

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("w:anchor"), href[1:])

        new_run = OxmlElement("w:r")

        rPr = OxmlElement("w:rPr")

        rColor = OxmlElement("w:color")
        rColor.set(qn("w:val"), "000080")
        rPr.append(rColor)

        rU = OxmlElement("w:u")
        rU.set(qn("w:val"), "single")
        rPr.append(rU)

        new_run.append(rPr)
        new_run.text = text

        hyperlink.append(new_run)

        if self.p:
            self.p._p.append(hyperlink)
        self.r = None

    def add_bookmark(self, anchor: str, text: str) -> None:
        if self.r:
            tag = self.r._r
            start = OxmlElement("w:bookmarkStart")
            start.set(qn("w:id"), "0")
            start.set(qn("w:name"), anchor)
            tag.addprevious(start)
            end = OxmlElement("w:bookmarkEnd")
            end.set(qn("w:id"), "0")
            tag.addnext(end)

            self.r.add_text(self.anchor + " " + text)

    def add_code(self, data: str) -> None:
        lines = data.splitlines()
        for linenr, line in enumerate(lines):
            self.add_text(line.rstrip())
            if linenr < len(lines) - 1:
                if self.r:
                    self.r.add_break()

    def add_list_style(self, name: str) -> None:
        self.finish_p()
        # The template included by python-docx only has 3 list styles.
        level = min(len(self.list_style) + 1, 3)
        suffix = f" {level}" if level > 1 else ""
        self.list_style.append(f"{name}{suffix}")

    def add_picture(self, attrs: List[Tuple[str, Optional[str]]]) -> None:
        src = get_attr(attrs, "src")
        height_attr = get_attr(attrs, "height")
        width_attr = get_attr(attrs, "width")
        height_px = int(height_attr) if height_attr else None
        width_px = int(width_attr) if width_attr else None

        image_buffer = load_image(src)
        size = image_size(image_buffer, width_px, height_px)
        paragraph = self.doc.add_paragraph()
        if self.alignment is not None:
            paragraph.alignment = self.alignment
        run = paragraph.add_run()
        run.add_picture(image_buffer, **size)

    def handle_starttag(self, tag: str, attrs: List[Tuple[str, Optional[str]]]) -> None:
        if tag == "a":
            self.href = get_attr(attrs, "href")
            self.anchor = get_attr(attrs, "id")
            self.init_run([])
        elif tag in ["b", "strong"]:
            self.init_run([("bold", True)])
        elif tag == "br":
            if self.r:
                self.r.add_break()
        elif tag == "code":
            self.init_run([("name", "Mono")])
        elif tag in ["em", "i"]:
            self.init_run([("italic", True)])
        elif tag in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            level = int(tag[-1])
            self.p = self.doc.add_heading(level=level)
        elif tag == "img":
            self.add_picture(attrs)
        elif tag == "ol":
            self.add_list_style("List Number")
        elif tag == "p":
            self.init_p(attrs)
        elif tag == "pre":
            self.pre = True
        elif tag == "style":
            self.style = True
        elif tag == "code":
            self.code = True
        elif tag == "span":
            span_attrs = html_attrs_to_font_style(attrs)
            self.init_run(span_attrs)
        elif tag == "sub":
            self.init_run([("subscript", True)])
        elif tag == "sup":
            self.init_run([("superscript", True)])
        elif tag == "u":
            self.init_run([("underline", True)])
        elif tag == "ul":
            self.add_list_style("List Bullet")
        elif tag == "table":
            self.init_table(attrs)
        elif tag == "tr":
            self.init_tr()
        elif tag in ["td", "th"]:
            self.init_tdth()

    def handle_data(self, data: str) -> None:
        if self.style:
            return
        if not self.pre:
            data = re.sub(WHITESPACE_RE, " ", data)
        if self.collapse_space:
            data = data.lstrip()
        if data:
            self.collapse_space = data.endswith(" ")

            if self.code:
                self.add_code(data)
            else:
                self.add_text(data)

    def handle_endtag(self, tag: str) -> None:
        if tag in ["a", "b", "code", "em", "i", "span", "strong", "sub", "sup", "u"]:
            self.finish_run()
            if tag == "a":
                self.href = ""
                self.anchor = ""
        elif tag in ["h1", "h2", "h3", "h4", "h5", "h6", "li", "ol", "p", "pre", "ul"]:
            self.finish_p()
        elif tag in ["ol", "ul"]:
            del self.list_style[-1]
        elif tag == "pre":
            self.pre = False
        elif tag == "style":
            self.style = False
        elif tag == "code":
            self.code = False
        elif tag == "table":
            self.finish_table()
        elif tag in ["td", "th"]:
            self.table_cell = None
            self.p = None
            self.r = None
